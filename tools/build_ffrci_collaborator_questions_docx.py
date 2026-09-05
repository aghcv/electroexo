from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


REPO_ROOT = Path(__file__).resolve().parents[1]
SOURCE_PATH = REPO_ROOT / "docs" / "ffrci_collaborator_questions.md"
OUTPUT_PATH = REPO_ROOT / "docs" / "ffrci_collaborator_questions.docx"


def set_font(
    run, name: str, size: float | None = None, bold: bool | None = None
) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.color.rgb = RGBColor(0, 0, 0)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def set_cell_margins(
    cell, top: int = 80, start: int = 100, bottom: int = 80, end: int = 100
) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (
        ("top", top),
        ("start", start),
        ("bottom", bottom),
        ("end", end),
    ):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_page_field(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_font(run, "Aptos", 9)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, text, end])


def add_inline_runs(paragraph, text: str, *, bold: bool = False) -> None:
    parts = re.split(r"(`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_font(run, "Aptos Mono", 10.5, bold)
        else:
            run = paragraph.add_run(part)
            set_font(run, "Aptos", 11, bold)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.start_type = WD_SECTION.NEW_PAGE
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.68)
    section.left_margin = Inches(0.82)
    section.right_margin = Inches(0.82)

    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(0, 0, 0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.08

    title = doc.styles["Title"]
    title.font.name = "Aptos Display"
    title._element.rPr.rFonts.set(qn("w:ascii"), "Aptos Display")
    title._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos Display")
    title.font.size = Pt(23)
    title.font.bold = True
    title.font.color.rgb = RGBColor(0, 0, 0)
    title.paragraph_format.space_after = Pt(16)
    title.paragraph_format.keep_with_next = True
    title_ppr = title.element.get_or_add_pPr()
    title_border = title_ppr.find(qn("w:pBdr"))
    if title_border is not None:
        title_ppr.remove(title_border)

    heading = doc.styles["Heading 1"]
    heading.font.name = "Aptos Display"
    heading._element.rPr.rFonts.set(qn("w:ascii"), "Aptos Display")
    heading._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos Display")
    heading.font.size = Pt(15)
    heading.font.bold = True
    heading.font.color.rgb = RGBColor(0, 0, 0)
    heading.paragraph_format.space_before = Pt(13)
    heading.paragraph_format.space_after = Pt(6)
    heading.paragraph_format.keep_with_next = True
    heading.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    footer = section.footer
    footer.is_linked_to_previous = False
    footer_table = footer.add_table(rows=1, cols=2, width=Inches(6.86))
    footer_table.columns[0].width = Inches(5.4)
    footer_table.columns[1].width = Inches(1.46)
    left = footer_table.cell(0, 0)
    right = footer_table.cell(0, 1)
    for cell in (left, right):
        set_cell_margins(cell, top=40, start=0, bottom=0, end=0)
    left_p = left.paragraphs[0]
    left_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    left_run = left_p.add_run("Electro Exocytosis Model Calibration")
    set_font(left_run, "Aptos", 9)
    add_page_field(right.paragraphs[0])


def build_docx(source: Path, output: Path) -> None:
    doc = Document()
    configure_document(doc)
    lines = source.read_text(encoding="utf-8").splitlines()
    index = 0
    while index < len(lines):
        line = lines[index].rstrip()
        if not line:
            index += 1
            continue
        if line.startswith("# "):
            paragraph = doc.add_paragraph(style="Title")
            add_inline_runs(paragraph, line[2:], bold=True)
            paragraph_ppr = paragraph._p.get_or_add_pPr()
            paragraph_border = paragraph_ppr.find(qn("w:pBdr"))
            if paragraph_border is not None:
                paragraph_ppr.remove(paragraph_border)
        elif line.startswith("## "):
            paragraph = doc.add_paragraph(style="Heading 1")
            add_inline_runs(paragraph, line[3:], bold=True)
        elif re.match(r"^\d+\.\s", line):
            match = re.match(r"^(\d+)\.\s+(.*)$", line)
            assert match is not None
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0)
            paragraph.paragraph_format.first_line_indent = Inches(0)
            paragraph.paragraph_format.space_after = Pt(5)
            paragraph.paragraph_format.line_spacing = 1.12
            run = paragraph.add_run(
                f"{match.group(1)}. {match.group(2).replace('`', '')}"
            )
            set_font(run, "Aptos", 10.8)
        elif line.startswith("   - "):
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0.62)
            paragraph.paragraph_format.first_line_indent = Inches(-0.22)
            paragraph.paragraph_format.space_after = Pt(2)
            bullet = paragraph.add_run("• ")
            set_font(bullet, "Aptos", 11)
            add_inline_runs(paragraph, line[5:])
        elif line.startswith("- "):
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0.34)
            paragraph.paragraph_format.first_line_indent = Inches(-0.20)
            paragraph.paragraph_format.space_after = Pt(4)
            bullet = paragraph.add_run("• ")
            set_font(bullet, "Aptos", 11)
            add_inline_runs(paragraph, line[2:])
        else:
            paragraph = doc.add_paragraph()
            add_inline_runs(paragraph, line)
        index += 1

    core = doc.core_properties
    core.title = "Experimental Metadata Needed for Electro Exocytosis Model Calibration"
    core.subject = (
        "Collaborator questionnaire for FFRCI particle and RNA sequencing data"
    )
    core.author = "Electro Exocytosis Modeling Team"
    core.keywords = "electro exocytosis, nsPEF, extracellular vesicles, Exoid, RNA sequencing, calibration"
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


if __name__ == "__main__":
    build_docx(SOURCE_PATH, OUTPUT_PATH)
    print(f"Wrote {OUTPUT_PATH}")

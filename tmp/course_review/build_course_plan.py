#!/usr/bin/env python3
"""Build a polished course-planning guide for computational EV biogenesis."""

from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path("/Users/aghorban/code/electroexo")
OUT = ROOT / "output" / "documents" / "Computational_Systems_Biology_for_EV_Biogenesis_Course_Plan.docx"

FIG_FRAMEWORK = Path("/Users/aghorban/code/electro-exocytosis/figs/fig01.png")
FIG_STORYBOARD = Path("/Users/aghorban/code/electro-exocytosis/figs/multilayer_storyboard.png")

PHYS_ROOT = Path(
    "/Users/aghorban/Library/CloudStorage/OneDrive-Personal/Documents/04-Universities/"
    "MCW_MU/Courses/BIEN5700-SystemPhysiology"
)
CSB_ROOT = Path(
    "/Users/aghorban/Library/CloudStorage/OneDrive-Personal/Documents/04-Universities/"
    "MCW_MU/Courses/BIEN6931-TopicsInBiomedicalEngineering/end of the semester back up/"
    "BIEN 6931 701 Topics in Biomedical Engr - 5132018 - 1103 PM"
)
REPORT_ROOT = Path("/Users/aghorban/code/electro-exocytosis")
CODE_ROOT = ROOT

NAVY = "1F3A5F"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "111827"
MUTED = "5B6573"
LIGHT_BLUE = "E8EEF5"
LIGHTER_BLUE = "F4F6F9"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "D8DEE8"
WHITE = "FFFFFF"
GOLD = "8A6A18"
RED = "9B1C1C"

PAGE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGIN = {"top": 80, "bottom": 80, "start": 120, "end": 120}


def rgb(hex_value: str) -> RGBColor:
    return RGBColor.from_string(hex_value)


def set_run_font(
    run,
    *,
    name: str = "Calibri",
    size: float | None = None,
    color: str | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, **kwargs: int) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin in ("top", "start", "bottom", "end"):
        if margin in kwargs:
            node = tc_mar.find(qn(f"w:{margin}"))
            if node is None:
                node = OxmlElement(f"w:{margin}")
                tc_mar.append(node)
            node.set(qn("w:w"), str(kwargs[margin]))
            node.set(qn("w:type"), "dxa")


def set_table_borders(table, color: str = MID_GRAY, size: int = 4) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), str(size))
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def set_table_geometry(table, widths_dxa: list[int]) -> None:
    if sum(widths_dxa) != PAGE_WIDTH_DXA:
        raise ValueError(f"Table widths must sum to {PAGE_WIDTH_DXA}: {widths_dxa}")
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(PAGE_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[idx]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell, **CELL_MARGIN)


def set_keep_with_next(paragraph, value: bool = True) -> None:
    paragraph.paragraph_format.keep_with_next = value


def set_repeat_row_no_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = tr_pr.find(qn("w:cantSplit"))
    if cant_split is None:
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_run_font(run, size=9, color=MUTED)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def add_numbering_definition(doc: Document, *, bullet: bool) -> int:
    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(el.get(qn("w:abstractNumId")))
        for el in numbering.findall(qn("w:abstractNum"))
        if el.get(qn("w:abstractNumId")) is not None
    ]
    num_ids = [
        int(el.get(qn("w:numId")))
        for el in numbering.findall(qn("w:num"))
        if el.get(qn("w:numId")) is not None
    ]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)

    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet" if bullet else "decimal")
    lvl.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "•" if bullet else "%1.")
    lvl.append(lvl_text)
    suff = OxmlElement("w:suff")
    suff.set(qn("w:val"), "tab")
    lvl.append(suff)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "271")
    p_pr.append(ind)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "300")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.append(spacing)
    lvl.append(p_pr)
    if bullet:
        r_pr = OxmlElement("w:rPr")
        fonts = OxmlElement("w:rFonts")
        fonts.set(qn("w:ascii"), "Calibri")
        fonts.set(qn("w:hAnsi"), "Calibri")
        r_pr.append(fonts)
        lvl.append(r_pr)
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_num(paragraph, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num_id_el)


def add_list_item(doc: Document, text: str, num_id: int, *, bold_lead: str | None = None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    apply_num(p, num_id)
    if bold_lead and text.startswith(bold_lead):
        lead = p.add_run(bold_lead)
        set_run_font(lead, size=11, color=INK, bold=True)
        rest = p.add_run(text[len(bold_lead) :])
        set_run_font(rest, size=11, color=INK)
    else:
        run = p.add_run(text)
        set_run_font(run, size=11, color=INK)
    return p


def add_body(doc: Document, text: str, *, bold_lead: str | None = None, after: float = 6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.25
    if bold_lead and text.startswith(bold_lead):
        lead = p.add_run(bold_lead)
        set_run_font(lead, size=11, color=INK, bold=True)
        rest = p.add_run(text[len(bold_lead) :])
        set_run_font(rest, size=11, color=INK)
    else:
        run = p.add_run(text)
        set_run_font(run, size=11, color=INK)
    return p


def add_kicker(doc: Document, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(5)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text.upper())
    set_run_font(run, size=10, color=GOLD, bold=True)
    return p


def add_source_note(doc: Document, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(text)
    set_run_font(run, size=8.5, color=MUTED, italic=True)
    return p


def add_callout(doc: Document, label: str, text: str, *, fill: str = LIGHTER_BLUE, accent: str = BLUE):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [PAGE_WIDTH_DXA])
    set_table_borders(table, color=accent, size=5)
    set_repeat_table_header(table.rows[0])
    set_repeat_row_no_split(table.rows[0])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.15
    r1 = p.add_run(f"{label}: ")
    set_run_font(r1, size=10.5, color=accent, bold=True)
    r2 = p.add_run(text)
    set_run_font(r2, size=10.5, color=INK)
    add_source_note(doc, "")
    return table


def add_table(
    doc: Document,
    headers: list[str],
    rows: list[list[str]],
    widths_dxa: list[int],
    *,
    font_size: float = 9.1,
    first_col_bold: bool = True,
) -> object:
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths_dxa)
    set_table_borders(table)
    set_repeat_table_header(table.rows[0])
    set_repeat_row_no_split(table.rows[0])
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.line_spacing = 1.05
        run = p.add_run(header)
        set_run_font(run, size=font_size, color=NAVY, bold=True)
    for row_values in rows:
        row = table.add_row()
        set_repeat_row_no_split(row)
        for idx, value in enumerate(row_values):
            cell = row.cells[idx]
            if len(table.rows) % 2 == 1:
                set_cell_shading(cell, "FAFBFC")
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.line_spacing = 1.08
            run = p.add_run(value)
            set_run_font(
                run,
                size=font_size,
                color=INK,
                bold=(first_col_bold and idx == 0),
            )
    set_table_geometry(table, widths_dxa)
    return table


def add_heading(doc: Document, text: str, level: int = 1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.add_run(text)
    set_keep_with_next(p)
    return p


def configure_styles(doc: Document) -> None:
    styles = doc.styles

    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_tokens = {
        1: (16, BLUE, 18, 10),
        2: (13, BLUE, 14, 7),
        3: (12, DARK_BLUE, 10, 5),
    }
    for level, (size, color, before, after) in heading_tokens.items():
        style = styles[f"Heading {level}"]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.0
        style.paragraph_format.keep_with_next = True

    if "Course Caption" not in styles:
        cap = styles.add_style("Course Caption", WD_STYLE_TYPE.PARAGRAPH)
    else:
        cap = styles["Course Caption"]
    cap.font.name = "Calibri"
    cap._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    cap._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    cap.font.size = Pt(9)
    cap.font.italic = True
    cap.font.color.rgb = rgb(MUTED)
    cap.paragraph_format.space_before = Pt(5)
    cap.paragraph_format.space_after = Pt(8)
    cap.paragraph_format.line_spacing = 1.0


def configure_section(section) -> None:
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.space_after = Pt(0)
    r1 = hp.add_run("COURSE PLANNING GUIDE")
    set_run_font(r1, size=8.5, color=BLUE, bold=True)
    r2 = hp.add_run("  |  Computational Systems Biology for EV Biogenesis")
    set_run_font(r2, size=8.5, color=MUTED)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.paragraph_format.space_before = Pt(0)
    fp.paragraph_format.space_after = Pt(0)
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_page_number(fp)

    first_header = section.first_page_header
    first_header.paragraphs[0].clear()
    first_footer = section.first_page_footer
    p = first_footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Curriculum development draft | August 2026")
    set_run_font(r, size=8.5, color=MUTED)


def add_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(40)
    p.paragraph_format.space_after = Pt(12)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("COURSE PLANNING GUIDE")
    set_run_font(run, size=10.5, color=GOLD, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(8)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("From Cell Physiology to Modular EV Biogenesis Simulation")
    set_run_font(run, size=27, color=NAVY, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(18)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("A beginner-first computational systems biology course built around electro-exocytosis")
    set_run_font(run, size=14.5, color=BLUE)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(18)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "Source review: BIEN5700 Systems Physiology, BIEN6931 Modeling and Simulations "
        "of Integrated Cellular Systems, the EV-biogenesis manuscript, and the executable electroexo model"
    )
    set_run_font(run, size=10.5, color=MUTED, italic=True)

    if FIG_FRAMEWORK.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(5)
        picture = p.add_run().add_picture(str(FIG_FRAMEWORK), width=Inches(6.15))
        picture._inline.docPr.set(
            "descr",
            "Integrated computational framework connecting nsPEF exposure, membrane electrodynamics, ion and calcium dynamics, repair, extracellular-vesicle biogenesis, cargo, injury, and manufacturing quality.",
        )
        picture._inline.docPr.set("title", "Integrated nsPEF-to-EV framework")
        cap = doc.add_paragraph(style="Course Caption")
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.add_run(
            "The course starts with familiar cell physiology and modeling patterns, then earns its way to this integrated nsPEF-to-EV framework."
        )

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    r = p.add_run("Prepared for long-range course development | August 2026")
    set_run_font(r, size=10, color=MUTED, bold=True)
    doc.add_page_break()


def build_document() -> Document:
    doc = Document()
    configure_styles(doc)
    configure_section(doc.sections[0])
    bullet_id = add_numbering_definition(doc, bullet=True)
    number_id = add_numbering_definition(doc, bullet=False)
    reading_number_id = add_numbering_definition(doc, bullet=False)

    add_cover(doc)

    add_heading(doc, "1. Executive recommendation", 1)
    add_callout(
        doc,
        "Recommended structure",
        "Use a 14-week beginner-first course for the full offering, and retain the existing 10-session slide sequence as a compressed seminar version. The full EV framework should appear only after students can explain a compartment, a mass balance, a membrane voltage, and an exposure-response chain in plain language.",
    )
    add_body(
        doc,
        "The reviewed materials support a coherent progression from physiology to modeling and then to the EV-biogenesis application. The strongest teaching spine is not the order of the manuscript. It is a repeated sequence: biological question, system boundary, stocks and flows, governing balance, simulation, and interpretation."
    )

    add_table(
        doc,
        ["Course block", "Weeks", "Purpose"],
        [
            ["I. Biological and modeling grammar", "1-7", "Homeostasis, transport, compartments, ODEs, membrane voltage, calcium, and rate laws."],
            ["II. Transferable modeling templates", "8-9", "PK/PD, PBPK/PD, physiological flows, variability, and model credibility."],
            ["III. EV-biogenesis capstone", "10-14", "Walk through the eight electroexo layers, integrate scenarios, and evaluate evidence and limits."],
        ],
        [2400, 900, 6060],
        font_size=9.6,
    )
    add_source_note(doc, "Course blocks are synthesized from the two course syllabi, the manuscript module map, and the current repository examples.")

    for item in [
        "Start with homeostasis and transport, not with code or the manuscript equations.",
        "Introduce ODEs through mass balance and visual stock-flow diagrams; analytical solution methods are optional enrichment.",
        "Use PK/PD and PBPK/PD as modeling templates for staged inputs, compartments, delayed effects, and physiological flows.",
        "Teach resting membrane potential and externally induced transmembrane voltage as related but distinct concepts.",
        "End every module with an input-state-output-evidence summary so students see how modules connect.",
    ]:
        add_list_item(doc, item, bullet_id)

    add_callout(
        doc,
        "Terminology note",
        "This plan interprets 'PBPD' as PBPK/PD. That matches the available PBPK_spring_2018.pdf and the standard progression from pharmacokinetics/pharmacodynamics to physiology-based models. If a different PBPD formulation was intended, it can be substituted in Week 9 without changing the course architecture.",
        fill="FFF8E8",
        accent=GOLD,
    )

    add_heading(doc, "2. Audience, outcomes, and teaching method", 1)
    add_heading(doc, "Audience and prerequisites", 2)
    add_body(
        doc,
        "The target audience is senior undergraduates, graduate trainees, clinicians, or life scientists who understand basic cell biology but have little or no computational background. Algebra, graph reading, and willingness to edit a few parameters are sufficient at entry; calculus and programming are introduced as tools rather than gatekeeping prerequisites."
    )

    add_heading(doc, "Learning outcomes", 2)
    outcomes = [
        "Define a biological system boundary and distinguish inputs, state variables, parameters, and outputs.",
        "Translate a stock-flow diagram into a mass-balance equation and explain every term biologically.",
        "Explain when a well-mixed compartment is useful and when spatial gradients or microdomains matter.",
        "Simulate simple ODE models and interpret time courses, steady states, thresholds, and sensitivity.",
        "Explain Nernst, GHK, membrane capacitance, channel gating, and induced membrane charging without requiring a full Hodgkin-Huxley derivation.",
        "Use PK/PD and PBPK/PD as analogies for staged exposure, distribution, effect, and physiological flow.",
        "Trace an nsPEF protocol through the eight EV-biogenesis model layers to yield, subtype, cargo, potency, purity, and viability.",
        "Critique calibration evidence, identify placeholder assumptions, and propose a focused validation or model-improvement project.",
    ]
    for item in outcomes:
        add_list_item(doc, item, bullet_id)

    add_heading(doc, "A repeated six-step lesson pattern", 2)
    lesson_steps = [
        "Begin with one biological question that can be answered in words.",
        "Draw the boundary and identify the compartments or pools.",
        "Label inflows, outflows, production, consumption, and feedback.",
        "Write the minimum governing equation and define units.",
        "Run or inspect one simulation, then perturb one parameter.",
        "Interpret the figure, name the assumption, and identify what evidence would test it.",
    ]
    for item in lesson_steps:
        add_list_item(doc, item, number_id)

    add_heading(doc, "3. Source review and chapter map", 1)
    add_body(
        doc,
        "The source folders contain substantial duplication from backups. The tables below identify the canonical lecture blocks that are most useful for this course and the page ranges that should be adapted. Page numbers refer to PDF pages in the named file, not textbook pages."
    )

    add_heading(doc, "BIEN5700 Systems Physiology: high-value background", 2)
    phys_rows = [
        ["Syllabus sp18.pdf", "p. 2", "Use the lesson-to-Guyton chapter map: Homeostasis ch. 1; membrane transport ch. 4; membrane potential ch. 5; muscle and neurotransmission chs. 6-8; hemodynamics ch. 14; exchange chs. 15-16; kidney chs. 25-29; metabolism chs. 68-70; energetics chs. 73-74; hormones ch. 75."],
        ["01homeostasis.pdf", "pp. 6-12", "Core: homeostasis, dynamic steady state, negative and positive feedback."],
        ["01homeostasis.pdf", "pp. 13-41", "Core: membrane structure, diffusion, electrical and pressure gradients, facilitated diffusion, osmosis, active transport, endocytosis, and exocytosis."],
        ["01homeostasis.pdf", "pp. 42-59", "Core: Nernst and GHK intuition, Na/K pump, action potential, voltage clamp, and patch clamp. Pages 60-65 are optional conduction examples."],
        ["02muscle.pdf", "pp. 33-48", "Support: neurotransmitter vesicles, depolarization, Ca2+ release, excitation-contraction coupling, and ATP-dependent recovery."],
        ["03circ.pdf", "pp. 18-38", "Support for PBPK/PD: pressure-flow relations, vascular compartments, capillary exchange, and transport."],
        ["04kidney.pdf", "pp. 2-4; 40-41", "Support for compartment volumes and clearance. The detailed nephron sequence is optional."],
        ["06respir.pdf", "pp. 30-45", "Optional case study for diffusion, exchange surfaces, transport capacity, and coupled compartments."],
        ["10metab.pdf", "pp. 1-10; 17-24", "Support for glycolysis, citric acid cycle, electron transport, ATP, energetics, and physiological stress."],
        ["11endo.pdf", "pp. 1-7; 30; 44", "Support for receptor signaling, second-messenger amplification, and feedback control."],
    ]
    add_table(doc, ["Source", "Pages", "Recommended teaching use"], phys_rows, [2450, 1050, 5860], font_size=8.65)
    add_source_note(doc, f"Reviewed root: {PHYS_ROOT}")

    add_heading(doc, "BIEN6931 Computational Systems Biology: core modeling sequence", 2)
    csb_rows_1 = [
        ["BIEN_6391_Spring_2018_Syllabus1.pdf", "p. 2", "Use modules 1-5 and 7 as the conceptual backbone. Cell-cycle modeling is a possible elective, not a prerequisite for EV biogenesis."],
        ["Dash_Lect1_Intro_Sim_Biol_Syst_Spring_2018.pdf", "pp. 2-6; 27-37", "Core: systems/computational biology, model purpose, iterative modeling, formulation, parameterization, validation, sensitivity, and communication."],
        ["Introduction to Simulation of Biological Systems.pdf", "pp. 1-20", "Core case study: model-building workflow and mechanistic versus nonmechanistic modeling through the aquarium example."],
        ["Introduction to Compartmental Modeling (1).pdf", "pp. 1-11", "Core: well-mixed assumption, amount-concentration-volume relations, one-compartment washout, reactions with flow, and two-compartment exchange."],
        ["Compartmental_modeling_supplement_spring_2018.pdf", "pp. 1-15", "Support: transport mechanisms, membrane fluxes, drug absorption/elimination, and hemodialysis."],
        ["Differential equations_Spring_2018.pdf", "pp. 1-6", "Core ODE vocabulary and initial conditions. Laplace transform pages 7-12 are optional enrichment."],
        ["Numerical_methods_Spring_2018.pdf", "pp. 1-11; 15-18", "Core: Euler, Runge-Kutta intuition, errors, convergence, and MATLAB ODE solvers. Stability pages 12-14 are instructor-level or advanced."],
        ["Introduction to Law of Mass Action.pdf", "pp. 1-5", "Core rate-law vocabulary and units for reversible and irreversible reactions."],
    ]
    add_table(doc, ["Source", "Pages", "Recommended teaching use"], csb_rows_1, [2750, 1200, 5410], font_size=8.55)
    add_source_note(doc, f"Reviewed canonical archive root: {CSB_ROOT}")

    add_heading(doc, "BIEN6931: application templates and advanced bridge", 2)
    csb_rows_2 = [
        ["Compartmental_Kinetic_PKPD_Modeling.pdf", "pp. 2-12; 24-29; 33-55", "Core: dose-to-effect stages, direct/inverse problems, one- and two-compartment models, absorption, half-life, volume of distribution, and nonlinear kinetics. Pages 13-15 add population variability."],
        ["PBPK_spring_2018.pdf", "pp. 1-15", "Core: physiology-based compartments, blood flow, permeability-surface area exchange, organ balances, and applications."],
        ["Facilitated_Diffusion.pdf", "pp. 1-6", "Core bridge from membrane biology to saturable carrier models."],
        ["Electrochemical_gradient_4_2_2018.pdf", "pp. 1-6", "Core: Fick, Nernst-Planck, GHK flux, and charged-solute compartment balances."],
        ["Resting_membrane_potential_4_2_2-18.pdf", "pp. 1-10", "Core: ion gradients, Nernst, resting potential, pump contribution, and voltage-gated channels."],
        ["Dash_Modeling Cellular Electrophysiology(1).pdf", "pp. 2-19", "Core/bridge: membranes, Nernst, current-voltage relations, gating variables, HH structure, GHK, and cardiac Ca2+ context. Pages 20-22 are advanced Markov-channel modeling."],
        ["Dash_The Hodgkin-Huxley Model.pdf", "pp. 1-10", "Optional advanced reading. Use for structure and gating intuition, not as a prerequisite derivation."],
        ["Dash_Modeling Biochemical Reaction Systems.pdf", "pp. 4-20", "Core: mass-conservation matrix, mass action, Gibbs energy, energy coupling, ion transport, enzymes, and Michaelis-Menten. Pages 21-31 are advanced reversible/cooperative/inhibition models."],
        ["Regression_analysis_handout_LM_Spring_2018.pdf", "pp. 1-8", "Guided advanced lab: nonlinear regression, local minima, Levenberg-Marquardt, parameter uncertainty, and example fits."],
        ["Model_assessment.pdf", "p. 1", "Core checklist for fit and parameter assessment."],
    ]
    add_table(doc, ["Source", "Pages", "Recommended teaching use"], csb_rows_2, [2760, 1240, 5360], font_size=8.35)
    add_source_note(doc, "The scanned handwritten thermodynamics, enzyme-kinetics, and voltage-gating notes are better retained as instructor references; the clean Dash lecture decks cover the same concepts more accessibly.")

    add_heading(doc, "Archive curation decisions", 2)
    curation_rows = [
        ["Use as the student spine", "01homeostasis; Intro Simulation; Intro Compartmental Modeling; ODEs/numerics; PKPD; PBPK; electrochemical gradient/resting potential; biochemical reactions; the electroexo examples."],
        ["Adapt selectively", "Muscle/neuro Ca2+ and exocytosis; circulation/kidney/respiratory case studies; full cellular electrophysiology; regression and model assessment."],
        ["Defer or make optional", "Laplace-transform derivations, exact eigenvalue solutions, full HH derivation, Markov-channel detail, cell-cycle module, whole-organ survey topics unrelated to the EV narrative."],
        ["Instructor-only references", "Handwritten scanned thermodynamics, simple-enzyme, and voltage-gating notes; backup ZIP duplicates; exam solutions and student submissions."],
    ]
    add_table(doc, ["Decision", "Materials"], curation_rows, [2100, 7260], font_size=9.0)
    add_source_note(doc, "This curation prevents the beginner course from inheriting the mathematical prerequisite load of the original graduate course.")

    add_heading(doc, "4. Recommended 14-week course sequence", 1)
    add_body(
        doc,
        "Each week assumes one concept-focused session plus one guided activity or lab. For a lecture-only format, the lab column becomes an in-class figure interpretation or parameter-prediction exercise."
    )

    schedule_1 = [
        ["1", "Why model living systems?", "Homeostasis, feedback, scales, system boundaries, models as useful simplifications.", "Draw a pulse-to-cell concept map before seeing the formal EV framework.", "BIEN5700 01homeostasis pp. 6-12; Dash Intro pp. 2-6, 27-37."],
        ["2", "Cells as compartments", "Membranes, organelles, transport, exocytosis, well-mixed assumptions, stocks and flows.", "Turn a generic cell into cytosol, ER, mitochondria, membrane, and extracellular compartments.", "01homeostasis pp. 13-41; Intro Compartmental pp. 1-3."],
        ["3", "Mass balance and ODEs", "Amount vs concentration, inflow/outflow, production/consumption, state variables, initial conditions.", "Build and interpret a one-compartment washout and a two-compartment exchange model.", "Intro Compartmental pp. 1-11; Differential Equations pp. 1-6."],
        ["4", "Numerical simulation", "Euler and Runge-Kutta intuition, solver output, error, time step, plots, reproducibility.", "Compare Euler with solve_ivp; explain when two curves disagree.", "Numerical Methods pp. 1-11, 15-18; Intro Simulation pp. 1-20."],
        ["5", "Membrane potential simulation", "Ion gradients, Nernst, GHK, pumps, capacitance, current versus voltage.", "Calculate single-ion equilibrium potentials, then vary permeabilities in a simple GHK model.", "01homeostasis pp. 42-59; Electrochemical Gradient pp. 1-6; Resting Potential pp. 1-10."],
        ["6", "Channels, calcium, and exocytosis", "Voltage gating, Ca2+ as signal and stressor, ER stores, synaptic release, patch-clamp interpretation.", "Annotate an action-potential/Ca2+ figure and identify which terms belong in a reduced ODE model.", "02muscle pp. 33-48; Cellular Electrophysiology pp. 2-19; HH optional."],
        ["7", "Biochemical rate laws", "Mass action, reversible reactions, saturation, Michaelis-Menten, Hill activation, energy coupling.", "Build a two-state reaction and a Hill switch; connect them to repair or secretion.", "Law of Mass Action pp. 1-5; Biochemical Reaction Systems pp. 4-20; Endocrine pp. 1-7."],
    ]
    add_table(
        doc,
        ["Wk", "Theme", "Core concepts", "Activity/lab", "Source anchors"],
        schedule_1,
        [560, 1470, 2390, 2320, 2620],
        font_size=7.95,
    )
    add_source_note(doc, "Weeks 1-7 establish the biological and mathematical grammar used repeatedly in the EV model.")

    schedule_2 = [
        ["8", "PK/PD as a template", "Dose-exposure-effect, one/two compartments, absorption, half-life, volume, nonlinear elimination, Emax/Hill effect.", "Simulate a one-compartment PK model and map concentration to effect.", "PKPD pp. 2-12, 24-29, 33-55."],
        ["9", "PBPK/PD and credibility", "Physiological compartments, blood flow, permeability, clearance, variability, fitting, sensitivity, and limits.", "Extend to two organs or compartments; perturb flow/permeability; write a model-credibility note.", "PBPK pp. 1-15; circulation pp. 18-38; kidney pp. 2-4, 40-41; Model Assessment."],
        ["10", "Layer 1: pulse and dosimetry", "Pulse amplitude, width, number, repetition rate, waveform, geometry, conductivity, energy, temperature.", "Run compare_dosimetry_models.py and separate field, dose, and heating interpretations.", "Manuscript Layer 1; README; compare_dosimetry_models.py."],
        ["11", "Layer 2: electrodynamics", "Resting Vm versus induced Delta Vm, RC/Schwan charging, organelle membranes, permeability and pore proxies.", "Run compare_membrane_electrodynamics.py; vary pulse width and cell radius.", "Manuscript Layer 2; electrodynamics.py; membrane-voltage figure."],
        ["12", "Layers 3-4: signaling and repair", "Ca2+ fluxes and stores, ROS, ATP, osmotic stress, PS, calpain, annexin, actin, resealing.", "Compare baseline, high-dose, Ca-limited, and repair-supported scenarios.", "ion_transport.py; remodeling_repair.py; comparison examples and figures."],
        ["13", "Layers 5-8: EV product", "MVB/ILV pools, fusion, budding, apoptotic bodies, cargo, potency, injury mixture, purity, recovery, batch effects.", "Contrast secretory, direct-loading, and injury-dominant regimes; defend which product is acceptable.", "ev_release.py; cargo_potency.py; injury_quality.py; manufacturing_qc.py."],
        ["14", "Integration and capstone", "Cell-state modifiers, multiscale storyboard, solution spaces, calibration order, evidence, and honest uncertainty.", "Present one submodel improvement, evidence table, validation plan, or teaching simulation.", "storyboard_multilayer_overview.py; solution_space_analysis; model assumptions and evidence targets."],
    ]
    add_table(
        doc,
        ["Wk", "Theme", "Core concepts", "Activity/lab", "Source anchors"],
        schedule_2,
        [560, 1470, 2390, 2320, 2620],
        font_size=7.9,
    )
    add_source_note(doc, "If a 15th week is available, split Week 9 into separate PBPK/PD and model-credibility weeks.")

    add_heading(doc, "5. From prerequisites to the modular EV-biogenesis framework", 1)
    add_callout(
        doc,
        "Plain-language organizing question",
        "What happens between an electrical pulse delivered outside a cell and a useful, high-quality EV product collected hours later? Each layer answers one part of that question and passes a compact set of outputs to the next layer.",
    )

    bridge_rows = [
        ["1. Pulse delivery and dosimetry", "Defines the external recipe and translates it into field, dose, energy, and temperature descriptors.", "Prerequisite: units, external inputs, energy balances. Demo: compare dosimetry formulations."],
        ["2. Membrane/organelle electrodynamics", "Converts the field into induced transmembrane voltages, conductance, permeability, and pore proxies.", "Prerequisite: membrane potential, capacitance, Nernst/GHK. Demo: pulse-width and cell-size sensitivity."],
        ["3. Ion transport, Ca2+, ROS, ATP", "Uses compartmental flux-balance ODEs to turn a brief electrical event into longer cellular trajectories.", "Prerequisite: compartments, fluxes, pumps, channels, reactions. Demo: baseline vs high dose vs Ca-limited."],
        ["4. Remodeling and repair", "Maps Ca2+ and stress into PS exposure, calpain, annexin, actin remodeling, lysosomal repair, and resealing.", "Prerequisite: Hill activation, feedback, competing processes. Demo: buffering, inhibition, repair support."],
        ["5. EV biogenesis and subtype release", "Tracks MVB/ILV/docking pools, plasma-membrane budding, fusion, and apoptotic-body production.", "Prerequisite: pool balances, exocytosis, rate laws. Demo: Rab/ESCRT-limited vs ceramide-favored vs injury-shifted."],
        ["6. Cargo, composition, and potency", "Combines subtype-weighted protein, RNA, lipid, antigen, and direct-loading states into functional output.", "Prerequisite: PK/PD-style effect mapping and saturation. Demo: endogenous sorting vs direct loading."],
        ["7. Injury, debris, and quality gate", "Separates useful EVs from stressed cells, apoptosis, necrosis, debris, and aggregates using thresholds and mixtures.", "Prerequisite: nonlinear switches, mixture accounting, model credibility. Demo: high particle count that fails quality."],
        ["8. Manufacturing, isolation, and QC", "Converts cellular release into isolated yield, recovery, purity, batch consistency, scale, and an optimization objective.", "Prerequisite: process mass balance, PBPK/flow intuition, constraints. Demo: yield-purity-viability tradeoff."],
        ["Cross-cutting cell-state modifiers", "Scales parameters for cell type, disease state, membrane composition, Ca2+ handling, metabolism, and stress susceptibility.", "Prerequisite: variability, parameter distributions, sensitivity. Demo: resilient-secretory vs fragile-stress-biased cells."],
    ]
    add_table(
        doc,
        ["Module", "What the module contributes", "Background concept and classroom demonstration"],
        bridge_rows,
        [2100, 3600, 3660],
        font_size=8.3,
    )
    add_source_note(doc, "Module names and interfaces follow main.tex, README.md, docs/model_assumptions.md, and the model files under electro_exocytosis/models.")

    add_heading(doc, "The module contract students should use", 2)
    for item in [
        "Input: what arrives from the experiment or the previous module?",
        "State: what quantities must be remembered over time?",
        "Process: what balance, rate law, threshold, or mapping changes those states?",
        "Output: what is passed downstream and what can be measured?",
        "Evidence: which experiment or source can constrain the parameters or test the behavior?",
        "Limitation: which assumption is most likely to fail for this cell type, timescale, or protocol?",
    ]:
        add_list_item(doc, item, bullet_id, bold_lead=item.split(":")[0] + ":")

    if FIG_STORYBOARD.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        picture = p.add_run().add_picture(str(FIG_STORYBOARD), width=Inches(6.35))
        picture._inline.docPr.set(
            "descr",
            "Three-row course storyboard comparing mild, productive, and injury-prone stimulation across exposure, electrical response, cell state, EV release, and quality outcomes.",
        )
        picture._inline.docPr.set("title", "EV biogenesis scenario storyboard")
        cap = doc.add_paragraph(style="Course Caption")
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.add_run(
            "Integrated storyboard across Layers 1-5. Use it after the individual modules, not before, so students can explain why mild, productive-secretory, and injury-dominant regimes diverge."
        )
        add_source_note(doc, f"Source figure: {FIG_STORYBOARD}")

    add_heading(doc, "6. Labs, assessment, and capstone", 1)
    add_heading(doc, "Computational scaffolding", 2)
    add_body(
        doc,
        "Use a no-code to guided-code to independent-analysis progression. In the first weeks, students predict curves and edit diagrams. By Weeks 3-5, they modify one equation or parameter in a prepared notebook. In the EV block, they run existing repository examples and explain the output before changing implementation details."
    )

    lab_items = [
        "Lab 1 - One-box washout: define volume, concentration, flow, time constant, and conserved amount.",
        "Lab 2 - Two-compartment exchange: change volume and permeability-surface area; explain equilibrium and timescale.",
        "Lab 3 - Membrane potential: compare Nernst single-ion predictions with a multi-ion GHK result.",
        "Lab 4 - PK/PD: connect a one-compartment concentration curve to an Emax or Hill effect curve.",
        "Lab 5 - PBPK/PD: add flow and tissue exchange; perform a one-parameter sensitivity analysis.",
        "Lab 6 - electroexo Layers 1-2: compare dosimetry and membrane charging across pulse protocols.",
        "Lab 7 - Layers 3-4: interpret Ca2+/ROS/ATP and repair trajectories across biological hypotheses.",
        "Lab 8 - Layers 5-8: classify a scenario as productive, direct-loading, or injury-dominant and justify the quality gate.",
    ]
    for item in lab_items:
        add_list_item(doc, item, bullet_id, bold_lead=item.split(":")[0] + ":")

    add_heading(doc, "Example assessment mix", 2)
    for item in [
        "15% concept checks and figure explanations",
        "35% guided computational labs",
        "15% source/evidence critique",
        "35% capstone proposal, simulation, presentation, and reflection",
    ]:
        add_list_item(doc, item, bullet_id)

    add_heading(doc, "Strong capstone formats", 2)
    for item in [
        "Improve one submodel and explain the biological assumption it replaces.",
        "Build one evidence-to-parameter table for a chosen layer and cell type.",
        "Design one validation experiment with measurable module inputs and outputs.",
        "Create one beginner-facing simulation or visualization that teaches a difficult concept correctly.",
        "Compare two cell-state parameterizations and identify which data would discriminate them.",
    ]:
        add_list_item(doc, item, bullet_id)

    add_callout(
        doc,
        "Scientific honesty requirement",
        "The current software is structurally complete but scientifically provisional. Students must distinguish a runnable model from a validated model and label placeholder parameters, assumptions, and evidence gaps explicitly.",
        fill="FFF1F1",
        accent=RED,
    )

    add_heading(doc, "7. Likely teaching pitfalls and how to avoid them", 1)
    pitfalls = [
        ["Starting with the full HH model", "Use Nernst, GHK, capacitance, and one gating variable first. HH is optional structure, not the entry point."],
        ["Confusing resting Vm with induced Delta Vm", "Name the source of each voltage every time: ionic gradients/permeability versus external-field charging."],
        ["Treating compartments as literal anatomy", "Repeat that a compartment is a modeling decision tied to mixing and timescale assumptions."],
        ["Using PK/PD as a biological equivalence", "Frame it as a reusable modeling pattern: input, disposition, state, and effect."],
        ["Stopping at total EV count", "Require subtype, cargo/potency, viability, contamination, and process recovery in the final interpretation."],
        ["Hiding placeholder status", "Pair every simulation with an assumptions/evidence box and a proposed validation measurement."],
        ["Showing the full framework too early", "Preview the destination briefly, then revisit it only after the prerequisite vocabulary has been built."],
    ]
    add_table(doc, ["Pitfall", "Teaching response"], pitfalls, [3000, 6360], font_size=9.0)
    add_source_note(doc, "These responses are consistent with the report's modular scope and the repository's explicit placeholder disclaimer.")

    add_heading(doc, "8. Existing assets to reuse", 1)
    asset_rows = [
        ["outputs/course_authoring/generate_electroexo_course_general_audience.mjs", "A 10-session physiology-first slide sequence. Retain as the compressed seminar version."],
        ["outputs/course_authoring/generate_electroexo_syllabus_ag_template.mjs", "A 14-week syllabus scaffold. Update it with the page-level source assignments in this guide."],
        ["outputs/course_authoring/generate_electroexo_course_slides.mjs", "A more technical 10-session deck. Use after the general-audience concepts or as graduate enrichment."],
        ["examples/compare_*.py", "Ready-made module demonstrations for dosimetry, electrodynamics, ion transport, repair, EV release, and downstream engineering."],
        ["examples/storyboard_multilayer_overview.py", "Best integrated visual narrative after students have learned the individual modules."],
        ["examples/solution_space_analysis", "Advanced exploration of pulse, waveform, conductivity, dosimetry, and cell-state parameter sweeps."],
        ["docs/model_assumptions.md and evidence targets", "Required reading for limitations, calibration order, and capstone design."],
    ]
    add_table(doc, ["Asset", "Recommended role"], asset_rows, [4550, 4810], font_size=8.65)
    add_source_note(doc, f"Executable repository reviewed: {CODE_ROOT}")

    add_heading(doc, "9. Production roadmap for next year's course", 1)
    roadmap = [
        ["T-9 to T-7 months", "Confirm audience, contact hours, software policy, and whether the primary format is 14 weeks or the 10-session seminar."],
        ["T-7 to T-5 months", "Create a canonical source folder, remove duplicates, and prepare a one-page reading guide for each week."],
        ["T-6 to T-4 months", "Build five starter notebooks: compartments, membrane potential, PK/PD, PBPK/PD, and electroexo module runner."],
        ["T-4 to T-3 months", "Revise the 14-week slide scaffold with the exact source pages and a common input-state-output-evidence visual language."],
        ["T-3 to T-2 months", "Pilot Weeks 3, 5, and 12 with noncomputational learners; measure where vocabulary or plots become confusing."],
        ["T-2 to T-1 months", "Finalize assessments, answer keys, accessibility checks, setup instructions, and a fallback no-install environment."],
        ["During delivery", "Collect one-minute reflections on what each equation term means biologically and revise the next offering accordingly."],
    ]
    add_table(doc, ["Timing", "Deliverable"], roadmap, [2100, 7260], font_size=9.1)
    add_source_note(doc, "The roadmap is intentionally semester-neutral so it can be aligned to the actual 2027 offering date.")

    add_heading(doc, "10. Canonical source roots and reading spine", 1)
    add_body(doc, "Use these roots as the stable source-of-truth locations while assembling the course:")
    roots = [
        f"Systems physiology: {PHYS_ROOT}",
        f"Computational systems biology: {CSB_ROOT}",
        f"EV-biogenesis manuscript and figures: {REPORT_ROOT}",
        f"Executable electroexo model and examples: {CODE_ROOT}",
    ]
    for item in roots:
        add_list_item(doc, item, bullet_id)

    add_heading(doc, "Minimal student reading spine", 2)
    spine = [
        "01homeostasis.pdf - homeostasis, membrane transport, exocytosis, and membrane electrical potentials.",
        "Dash_Lect1_Intro_Sim_Biol_Syst_Spring_2018.pdf - what computational systems biology is and how models are developed.",
        "Introduction to Compartmental Modeling (1).pdf - the core mass-balance examples.",
        "Numerical_methods_Spring_2018.pdf - solver intuition and numerical error.",
        "Compartmental_Kinetic_PKPD_Modeling.pdf and PBPK_spring_2018.pdf - staged exposure/effect and physiology-based compartments.",
        "Electrochemical_gradient_4_2_2018.pdf and Resting_membrane_potential_4_2_2-18.pdf - charged transport and voltage.",
        "Dash_Modeling Biochemical Reaction Systems.pdf - rate laws, energy coupling, and enzyme kinetics.",
        "The manuscript module map plus README.md and docs/model_assumptions.md - the final EV framework and its limitations.",
    ]
    for item in spine:
        add_list_item(doc, item, reading_number_id)

    add_heading(doc, "Immediate next build", 2)
    add_callout(
        doc,
        "Recommended next artifact",
        "Update the existing 14-week syllabus and slide authoring pipeline with these exact page assignments, then build the first four starter notebooks. That sequence locks the course architecture before substantial slide production begins.",
    )
    add_heading(doc, "Planning decisions to confirm", 2)
    decisions = [
        "Primary audience, assumed biology background, and whether any coding experience can be expected.",
        "Full 14-week course, compressed 10-session seminar, or both as linked formats.",
        "Contact time, laboratory time, and the balance between discussion, equations, and hands-on simulation.",
        "A no-install computational environment, accessibility requirements, and the level of Python exposed to students.",
        "Capstone team size, deliverables, evidence expectations, and the boundary between explanation and model extension.",
        "One recurring cell type and EV application that can anchor examples across the semester.",
        "Language for distinguishing validated mechanisms, plausible hypotheses, and exploratory model components.",
    ]
    for item in decisions:
        add_list_item(doc, item, bullet_id)
    add_callout(
        doc,
        "Course-design filter",
        "Keep a concept, equation, or computational exercise only when it helps students explain a model interface, interpret an output, or judge the strength of the supporting evidence.",
    )

    return doc


def main() -> None:
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = build_document()
    doc.core_properties.title = "From Cell Physiology to Modular EV Biogenesis Simulation"
    doc.core_properties.subject = "Course planning guide for computational systems biology and electro-exocytosis"
    doc.core_properties.author = "Course planning draft"
    doc.core_properties.keywords = "computational systems biology, EV biogenesis, electro-exocytosis, compartmental modeling, PKPD, PBPK"
    doc.core_properties.comments = "Prepared from local course archives, manuscript materials, and executable model assets."
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
